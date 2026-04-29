import streamlit as st
import cnlunar
import datetime
import io
import os
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches

# ==================== Constants ====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_IMAGE_PATH = os.path.join(SCRIPT_DIR, "yellow.jpg")
FONT_PATH = "C:/Windows/Fonts/kaiu.ttf"
SCALE = 5  # Upscale factor for image quality


# ==================== 八字 Calculation ====================
def calculate_bazi(birth_dt):
    """Calculate 八字 using cnlunar library."""
    lunar = cnlunar.Lunar(birth_dt)
    return {
        "year": lunar.year8Char,
        "month": lunar.month8Char,
        "day": lunar.day8Char,
        "hour": lunar.twohour8Char,
    }


# ==================== Image Generation ====================
def draw_vertical_text(draw, text, center_x, start_y, font, fill, spacing=0):
    """Draw text vertically (top to bottom), centered horizontally on center_x.
    Returns the y position after the last character.
    """
    y = start_y
    for char in text:
        bbox = draw.textbbox((0, 0), char, font=font)
        cw = bbox[2] - bbox[0]
        ch = bbox[3] - bbox[1]
        x = center_x - cw / 2
        draw.text((x, y), char, font=font, fill=fill)
        y += ch + spacing
    return y


def create_bazi_image(client_name, bazi):
    """Create the 八字 card image with vertical text overlay."""
    # Load and upscale the base image for better text quality
    base = Image.open(BASE_IMAGE_PATH).convert("RGB")
    w, h = base.size[0] * SCALE, base.size[1] * SCALE
    img = base.resize((w, h), Image.LANCZOS)
    draw = ImageDraw.Draw(img)

    # Font sizes (proportional to image width)
    title_size = int(w * 0.09)
    main_size = int(w * 0.07)
    label_size = int(w * 0.05)

    title_font = ImageFont.truetype(FONT_PATH, title_size)
    main_font = ImageFont.truetype(FONT_PATH, main_size)
    label_font = ImageFont.truetype(FONT_PATH, label_size)

    text_color = (139, 0, 0)       # Dark red
    label_color = (110, 20, 20)    # Slightly lighter

    # --- Title (centered, near top) ---
    title = "八字命盤"
    tbbox = draw.textbbox((0, 0), title, font=title_font)
    tw = tbbox[2] - tbbox[0]
    draw.text(((w - tw) / 2, h * 0.04), title, font=title_font, fill=text_color)

    # --- Vertical columns (right to left: name, 年, 月, 日, 時) ---
    start_y = h * 0.15
    char_spacing = int(main_size * 0.35)
    label_gap = int(main_size * 0.5)

    # Column x positions (proportional, right to left)
    name_x = w * 0.84
    year_x = w * 0.66
    month_x = w * 0.50
    day_x = w * 0.34
    hour_x = w * 0.18

    # Draw client name vertically
    draw_vertical_text(
        draw, client_name, name_x, start_y,
        main_font, text_color, spacing=char_spacing,
    )

    # Draw four pillars
    pillars = [
        ("年", bazi["year"], year_x),
        ("月", bazi["month"], month_x),
        ("日", bazi["day"], day_x),
        ("時", bazi["hour"], hour_x),
    ]

    for label, chars, col_x in pillars:
        # Draw label (年/月/日/時)
        y = start_y
        y = draw_vertical_text(
            draw, label, col_x, y,
            label_font, label_color, spacing=0,
        )
        y += label_gap
        # Draw stem and branch characters
        draw_vertical_text(
            draw, chars, col_x, y,
            main_font, text_color, spacing=char_spacing,
        )

    return img


# ==================== Word Document ====================
def create_word_doc(client_name, bazi, birth_date, birth_time, img):
    """Create a Word document containing the 八字 info and card image."""
    doc = Document()
    doc.add_heading(f"{client_name} 八字命盤", level=1)
    doc.add_paragraph(f"出生日期：{birth_date.strftime('%Y年%m月%d日')}")
    doc.add_paragraph(f"出生時間：{birth_time.strftime('%H:%M')}")
    doc.add_paragraph("")

    # 八字 table
    table = doc.add_table(rows=3, cols=4, style="Table Grid")
    for i, label in enumerate(["年柱", "月柱", "日柱", "時柱"]):
        table.rows[0].cells[i].text = label
    for i, key in enumerate(["year", "month", "day", "hour"]):
        table.rows[1].cells[i].text = bazi[key][0]   # Heavenly Stem
        table.rows[2].cells[i].text = bazi[key][1]   # Earthly Branch

    doc.add_paragraph("")

    # Insert the card image
    img_buf = io.BytesIO()
    img.save(img_buf, format="PNG")
    img_buf.seek(0)
    doc.add_picture(img_buf, width=Inches(3))

    # Save document to buffer
    doc_buf = io.BytesIO()
    doc.save(doc_buf)
    doc_buf.seek(0)
    return doc_buf


# ==================== Streamlit App ====================
st.set_page_config(page_title="八字命盤產生器", layout="centered")
st.title("八字命盤產生器")

client_name = st.text_input("客戶姓名", placeholder="請輸入客戶姓名")

col1, col2 = st.columns(2)
with col1:
    birth_date = st.date_input(
        "出生日期（國曆）",
        value=datetime.date(1990, 1, 1),
        min_value=datetime.date(1900, 1, 1),
        max_value=datetime.date.today(),
    )
with col2:
    birth_time = st.time_input("出生時間", value=datetime.time(12, 0))

if st.button("生成八字", type="primary", use_container_width=True):
    if not client_name.strip():
        st.error("請輸入客戶姓名")
    else:
        birth_dt = datetime.datetime.combine(birth_date, birth_time)
        bazi = calculate_bazi(birth_dt)
        st.session_state["bazi"] = bazi
        st.session_state["client_name"] = client_name.strip()
        st.session_state["birth_date"] = birth_date
        st.session_state["birth_time"] = birth_time

# Display results if available
if "bazi" in st.session_state:
    bazi = st.session_state["bazi"]
    name = st.session_state["client_name"]

    st.divider()
    st.subheader(f"{name} 的八字")

    # Show four pillars as metrics
    cols = st.columns(4)
    for col, (label, key) in zip(
        cols, [("年柱", "year"), ("月柱", "month"), ("日柱", "day"), ("時柱", "hour")]
    ):
        with col:
            st.metric(label, bazi[key])

    # Generate card image
    img = create_bazi_image(name, bazi)
    st.image(img, caption=f"{name} 八字命盤", use_container_width=True)

    # Generate Word document and show download button
    doc_buf = create_word_doc(
        name, bazi,
        st.session_state["birth_date"],
        st.session_state["birth_time"],
        img,
    )
    st.download_button(
        label="下載 Word 檔案",
        data=doc_buf,
        file_name=f"{name}_八字命盤.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
